import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinterdnd2 import DND_FILES, TkinterDnD
from PyPDF2 import PdfReader
import os
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl
from openpyxl import Workbook

def open_file():
    filepath = filedialog.askopenfilename(filetypes=[("PDF files","*.pdf"), ("All Files","*.*")])
    if filepath:
        selected_file.set(filepath)
        
def drop(event):
    file_path = event.data
    if file_path.startswith("{") and file_path.endswith("}"):
        file_path = file_path[1:-1]
    selected_file.set(file_path)

def choose_directory():
    directory = filedialog.askdirectory()
    output_dir.set(directory)




def start_conversion_pdf():
    try:
        pdf_path = selected_file.get()
        if not pdf_path:
            messagebox.showerror("Error", "PDF file not selected")
            return
        if not output_file.get():
            raise ValueError("Output file name cannot be empty")
        output_dir_path = output_dir.get()
        if not output_dir_path:
            messagebox.showerror("Error", "No output directory selected")
            return
        
        progress['value'] = 0
        output_path = os.path.join(output_dir_path, f"{output_file.get()}.txt")
        with open(pdf_path, "rb") as pdf_file:
            reader = PdfReader(pdf_file)
            total_pages = len(reader.pages)
            
            with open(output_path, "w", encoding="UTF-8") as text_file:
                for i, page in enumerate(reader.pages):
                    progress["value"] = ((i + 1) / total_pages) * 100
                    root.update_idletasks()
                    time.sleep(0.5)
                    text = page.extract_text()
                    
                    if text:
                        text_file.write(f"Page {i + 1}\n{'=' * 20}\n")
                        text_file.write(text)
                        text_file.write("\n\n")
        
        messagebox.showinfo("Success", "File converted successfully!")
        
    except Exception as e:
        messagebox.showerror("Error", str(e))
       
def start_conversion_docx():
    try:
        pdf_path = selected_file.get()
        if not pdf_path:
            messagebox.showerror("Error", "PDF file not selected")
            return
        if not output_file.get():
            raise ValueError("Output file name cannot be empty")
        output_dir_path = output_dir.get()
        if not output_dir_path:
            messagebox.showerror("Error", "No output directory selected")
            return
        
        progress['value'] = 0
        output_path = os.path.join(output_dir_path, f"{output_file.get()}.docx")
        with open(pdf_path, "rb") as pdf_file:
            reader = PdfReader(pdf_file)
            total_pages = len(reader.pages)
            
            doc = Document()
            for i, page in enumerate(reader.pages):
                progress["value"] = ((i + 1) / total_pages) * 100
                root.update_idletasks()
                time.sleep(0.5)
                text = page.extract_text()
                
                if text:
                    page_number_paragraph = doc.add_paragraph(f"Page {i + 1}")
                    page_number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    page_number_paragraph.bold = True
                    
                   
                    separator_paragraph = doc.add_paragraph("=" * 20)
                    separator_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = separator_paragraph.runs[0]
                    run.bold = True
                    
                   
                    main_text_paragraph = doc.add_paragraph(text)
                    main_text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    
                    
            doc.save(output_path)
        
        messagebox.showinfo("Success", "File converted successfully!")
        
    except Exception as e:
         messagebox.showerror("Error", str(e))


def start_conversion_xlsx():
    try:
        pdf_path = selected_file.get()
        if not pdf_path:
            messagebox.showerror("Error", "PDF file not selected")
            return
        if not output_file.get():
            raise ValueError("Output file name cannot be empty")
        output_dir_path = output_dir.get()
        if not output_dir_path:
            messagebox.showerror("Error", "No output directory selected")
            return
        
        progress['value'] = 0
        output_path = os.path.join(output_dir_path, f"{output_file.get()}.xlsx")
        try:
            with open(pdf_path, "rb") as pdf_file:
                reader = PdfReader(pdf_file)
                total_pages = len(reader.pages)
                
                workbook = Workbook()
                sheet = workbook.active
                row_index = 1  # Początkowy indeks wiersza w arkuszu
                
                for i, page in enumerate(reader.pages):
                    progress["value"] = ((i + 1) / total_pages) * 100
                    root.update_idletasks()
                    time.sleep(0.5)
                    text = page.extract_text()
                    
                    if text:
                        lines = text.splitlines()
                        for line in lines:
                            cleaned_line = line.strip()  
                            if cleaned_line:  
                                sheet.cell(row=row_index, column=1, value=cleaned_line)
                                row_index += 1
                    
                        
                        row_index += 1
                
                workbook.save(output_path)
                
            messagebox.showinfo("Success", "File converted successfully!")
            
        except Exception as internal_error:
           
            messagebox.showerror("Internal Error", f"An error occurred during the conversion process: {str(internal_error)}")
    
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

# MAIN WINDOW

root = TkinterDnD.Tk()
root.title("PDF 2 TXT converter")
root.geometry("500x400")

root.drop_target_register(DND_FILES)
root.dnd_bind("<<Drop>>", drop)

output_dir = tk.StringVar()
output_file = tk.StringVar()
selected_file = tk.StringVar()

tk.Label(root, text="Output directory").pack(pady=5) 
tk.Entry(root, textvariable=output_dir).pack(padx=10, pady=5)  
tk.Button(root, text="Browse", command=choose_directory).pack(pady=5)  

tk.Label(root, text="Output file name").pack(pady=5)  
tk.Entry(root, textvariable=output_file).pack(padx=10, pady=5)  

tk.Label(root, text="Select or drop PDF file").pack(pady=5) 
tk.Entry(root, textvariable=selected_file, state="readonly").pack(padx=10, pady=5)  

open_button = tk.Button(root, text="Open PDF", command=open_file)
open_button.pack(pady=5) 

convert_button_pdf = tk.Button(root, text="Convert txt", command=start_conversion_pdf)
convert_button_docx = tk.Button(root, text="Convert docx", command=start_conversion_docx)
#convert_button_xslx = tk.Button(root, text="Convert xlsx", command=start_conversion_xlsx)
convert_button_pdf.pack(pady=5) 
convert_button_docx.pack(padx=5, pady=0) 
#convert_button_xslx.pack(padx=5) # Zwiększony odstęp pionowy

progress = ttk.Progressbar(root, orient="horizontal", length=200, mode="determinate")
progress.pack(pady=10) 

root.mainloop()
